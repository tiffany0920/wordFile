import os
import logging
import markdown
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from typing import Optional
import re
from config import Config
from typing import Tuple
import hashlib
import requests

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class WordConverter:
    """Markdown到Word转换器类"""
    
    def __init__(self):
        """初始化Word转换器"""
        self.output_dir = Config.OUTPUT_DIR
        os.makedirs(self.output_dir, exist_ok=True)
    
    def markdown_to_word(self, markdown_content: str, output_filename: Optional[str] = None) -> str:
        """
        将Markdown内容转换为Word文档

        Args:
            markdown_content: Markdown内容
            output_filename: 输出文件名

        Returns:
            生成的Word文件路径
        """
        try:
            logger.info("开始转换Markdown到Word...")
            logger.info(f"输出文件名: {output_filename}")

            # 预处理：下载远程图片并重写为本地 media 路径
            markdown_content = self._preprocess_markdown_assets(markdown_content)
            logger.info("预处理完成，开始解析Markdown内容...")

            # 创建新的Word文档
            doc = Document()
            
            # 解析Markdown内容（含图片、表格块）
            lines = [ln.rstrip('\r') for ln in markdown_content.split('\n')]
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    doc.add_paragraph()
                    i += 1
                    continue
                # 图片: ![alt](path "title") 或 ![alt](path)
                img_match = re.match(r'^!\[(.*?)\]\(([^\)\s]+)(?:\s+"(.*?)")?\)$', line)
                if img_match:
                    alt_text = img_match.group(1)
                    img_path = img_match.group(2)
                    self._add_image(doc, img_path, alt_text)
                    i += 1
                    continue
                # 表格块：header | line, 分隔线 |---|---|，随后多行数据
                if self._is_table_header(line) and i + 1 < len(lines) and self._is_table_separator(lines[i + 1].strip()):
                    i = self._add_table_block(doc, lines, i)
                    continue
                if line.startswith('# '):
                    heading = doc.add_heading(line[2:], level=1)
                    self._format_heading(heading)
                elif line.startswith('## '):
                    heading = doc.add_heading(line[3:], level=2)
                    self._format_heading(heading)
                elif line.startswith('### '):
                    heading = doc.add_heading(line[4:], level=3)
                    self._format_heading(heading)
                elif line.startswith('#### '):
                    heading = doc.add_heading(line[5:], level=4)
                    self._format_heading(heading)
                elif line.startswith('- ') or line.startswith('* '):
                    self._add_list_item(doc, line[2:], is_ordered=False)
                elif re.match(r'^\d+\. ', line):
                    self._add_list_item(doc, re.sub(r'^\d+\. ', '', line), is_ordered=True)
                elif line.startswith('> '):
                    self._add_quote(doc, line[2:])
                elif line.startswith('```'):
                    # 简化处理：跳过代码围栏行，代码正文以普通段落加入
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        self._add_paragraph(doc, lines[i])
                        i += 1
                else:
                    self._add_paragraph(doc, line)
                i += 1
            
            # 生成输出文件名
            if not output_filename:
                from datetime import datetime
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"converted_document_{timestamp}.docx"
            elif not output_filename.endswith('.docx'):
                output_filename += '.docx'
            
            # 保存Word文档
            output_path = os.path.join(self.output_dir, output_filename)
            doc.save(output_path)
            
            logger.info(f"Word文档已保存: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"转换Markdown到Word时出错: {str(e)}")
            raise Exception(f"转换Markdown到Word失败: {str(e)}")

    def markdown_to_word_pandoc(self, markdown_content: str, output_filename: Optional[str] = None) -> str:
        """使用 Pandoc 高保真将 Markdown 转换为 Word。"""
        try:
            import pypandoc
            logger.info("使用 Pandoc 将 Markdown 转为 Word...")
            # 预处理：下载远程图片并重写为本地 media 路径
            markdown_content = self._preprocess_markdown_assets(markdown_content)
            if not output_filename:
                from datetime import datetime
                output_filename = f"converted_document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            elif not output_filename.endswith('.docx'):
                output_filename += '.docx'
            output_path = os.path.join(self.output_dir, output_filename)
            # 资源路径：输出目录与其 media 子目录，供图片/表格资源查找
            media_dir = os.path.join(self.output_dir, 'media')
            os.makedirs(media_dir, exist_ok=True)
            resource_paths = [self.output_dir, media_dir]
            extra_args = [
                f"--resource-path={os.pathsep.join(resource_paths)}",
            ]
            # Pandoc 转换（使用最稳定的格式）
            pypandoc.convert_text(
                markdown_content,
                'docx',
                format='markdown',
                outputfile=output_path,
                extra_args=extra_args,
            )
            logger.info(f"Pandoc 转换完成: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Pandoc 转换Markdown到Word失败: {str(e)}")
            # 回退到内置转换
            logger.info("回退到内置转换器")
            return self.markdown_to_word(markdown_content, output_filename)
    
    def _format_heading(self, heading):
        """格式化标题"""
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_list_item(self, doc, text, is_ordered=False):
        """添加列表项"""
        p = doc.add_paragraph()
        if is_ordered:
            p.style = 'List Number'
        else:
            p.style = 'List Bullet'
        p.add_run(text)
    
    def _add_quote(self, doc, text):
        """添加引用"""
        p = doc.add_paragraph()
        p.style = 'Quote'
        p.add_run(text)
    
    def _add_table_row(self, doc, line):
        """添加表格行"""
        # 简单的表格处理，实际应用中可能需要更复杂的逻辑
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if len(cells) > 1:
            table = doc.add_table(rows=1, cols=len(cells))
            table.style = 'Table Grid'
            for i, cell in enumerate(cells):
                table.cell(0, i).text = cell

    def _is_table_header(self, line: str) -> bool:
        return line.startswith('|') and '|' in line[1:]

    def _is_table_separator(self, line: str) -> bool:
        if not (line.startswith('|') and '|' in line[1:]):
            return False
        parts = [p.strip() for p in line.split('|')[1:-1]]
        if not parts:
            return False
        return all(re.match(r'^:?-{3,}:?$', p) for p in parts)

    def _add_table_block(self, doc, lines, start_idx: int) -> int:
        """从表头开始创建完整表格，返回消耗后的下一行索引"""
        header_line = lines[start_idx].strip()
        sep_idx = start_idx + 1
        headers = [c.strip() for c in header_line.split('|')[1:-1]]
        # 收集数据行
        data_rows = []
        i = sep_idx + 1
        while i < len(lines):
            row = lines[i].strip()
            if not row.startswith('|') or '|' not in row[1:]:
                break
            cells = [c.strip() for c in row.split('|')[1:-1]]
            if len(cells) != len(headers):
                break
            data_rows.append(cells)
            i += 1
        # 创建表格（1 header row + n data rows）
        cols = max(1, len(headers))
        rows = 1 + len(data_rows)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        # header
        for c_idx, h in enumerate(headers):
            table.cell(0, c_idx).text = h
        # data
        for r_idx, row_cells in enumerate(data_rows, start=1):
            for c_idx, cell in enumerate(row_cells):
                table.cell(r_idx, c_idx).text = cell
        return i

    def _add_image(self, doc, image_path: str, alt_text: str = ""):
        """插入图片，若找不到路径则显示占位符而不显示假链接"""
        try:
            logger.info(f"尝试插入图片: {image_path}, alt_text: {alt_text}")

            # 支持远程URL：已在预处理阶段下载到本地 media 目录
            if image_path.startswith('http://') or image_path.startswith('https://'):
                logger.info(f"处理远程URL图片: {image_path}")
                local_path = self._download_image_to_media(image_path)
                if local_path is None:
                    logger.warning(f"无法下载远程图片: {image_path}")
                    # 显示占位符而不是假链接
                    p = doc.add_paragraph()
                    run = p.add_run(f"[图片: {alt_text or '未命名图片'}]")
                    run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
                    return
                image_path = local_path
                logger.info(f"远程图片已下载到: {image_path}")

            # 支持相对路径（相对于输出目录）
            candidate_paths = [
                image_path,
                os.path.join(self.output_dir, image_path),
            ]
            actual = None
            for p in candidate_paths:
                logger.info(f"检查路径: {p}")
                if os.path.isfile(p):
                    actual = p
                    logger.info(f"找到图片文件: {actual}")
                    break

            # 如果还找不到，尝试从media目录查找
            if actual is None and not os.path.isabs(image_path):
                media_dir = os.path.join(self.output_dir, 'media')
                logger.info(f"在media目录中查找: {media_dir}")

                if image_path.startswith('media/'):
                    # 如果路径已经是 media/filename 格式
                    actual_path = os.path.join(self.output_dir, image_path)
                    logger.info(f"检查media格式路径: {actual_path}")
                    if os.path.isfile(actual_path):
                        actual = actual_path
                        logger.info(f"在media目录中找到图片: {actual}")
                    else:
                        # 尝试直接在media目录中查找文件名
                        filename = os.path.basename(image_path)
                        media_path = os.path.join(media_dir, filename)
                        logger.info(f"检查media目录中的文件名: {media_path}")
                        if os.path.isfile(media_path):
                            actual = media_path
                            logger.info(f"在media目录中找到图片: {actual}")
                else:
                    # 如果只是文件名，尝试在media目录中查找
                    media_path = os.path.join(media_dir, os.path.basename(image_path))
                    logger.info(f"检查media目录中的文件: {media_path}")
                    if os.path.isfile(media_path):
                        actual = media_path
                        logger.info(f"在media目录中找到图片: {actual}")

                # 最后尝试：检查嵌套的media/media目录
                if actual is None:
                    nested_media_path = os.path.join(media_dir, image_path)
                    logger.info(f"检查嵌套media目录: {nested_media_path}")
                    if os.path.isfile(nested_media_path):
                        actual = nested_media_path
                        logger.info(f"在嵌套media目录中找到图片: {actual}")

                        # 将文件复制到正确的media目录
                        correct_path = os.path.join(media_dir, os.path.basename(image_path))
                        if correct_path != nested_media_path:
                            import shutil
                            shutil.copy2(nested_media_path, correct_path)
                            logger.info(f"复制文件到正确位置: {nested_media_path} -> {correct_path}")
                            actual = correct_path

            if actual is None:
                logger.error(f"无法找到图片文件: {image_path}")
                # 显示占位符而不是假链接
                p = doc.add_paragraph()
                run = p.add_run(f"[图片: {alt_text or '未命名图片'}]")
                run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
                return

            logger.info(f"成功插入图片: {actual}")
            # 默认宽度 5 英寸，可扩展解析 title 指定尺寸
            doc.add_picture(actual, width=Inches(5))
            if alt_text:
                p = doc.add_paragraph(alt_text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            logger.error(f"插入图片失败: {str(e)}")
            # 显示占位符而不是假链接
            p = doc.add_paragraph()
            run = p.add_run(f"[图片加载失败: {alt_text or '未命名图片'}]")
            run.font.color.rgb = RGBColor(255, 0, 0)  # 红色

    def _preprocess_markdown_assets(self, markdown_content: str) -> str:
        """下载远程图片为本地 media 路径，并重写Markdown中的链接。"""
        os.makedirs(self.output_dir, exist_ok=True)
        media_dir = os.path.join(self.output_dir, 'media')
        os.makedirs(media_dir, exist_ok=True)

        pattern = re.compile(r'!\[(?P<alt>[^\]]*)\]\((?P<url>https?://[^\)\s]+)(?:\s+"(?P<title>.*?)")?\)')

        def repl(m: re.Match) -> str:
            url = m.group('url')
            alt = m.group('alt') or 'image'
            title = m.group('title')
            local = self._download_image_to_media(url)
            if local is None:
                return m.group(0)
            rel_path = os.path.relpath(local, self.output_dir).replace('\\', '/')
            if title:
                return f"![{alt}]({rel_path} \"{title}\")"
            return f"![{alt}]({rel_path})"

        return pattern.sub(repl, markdown_content)

    def _download_image_to_media(self, url: str) -> Optional[str]:
        try:
            os.makedirs(self.output_dir, exist_ok=True)
            media_dir = os.path.join(self.output_dir, 'media')
            os.makedirs(media_dir, exist_ok=True)
            # 生成稳定文件名
            ext = os.path.splitext(url.split('?')[0].split('#')[0])[1]
            if not ext or len(ext) > 5:
                ext = '.png'
            name = hashlib.sha256(url.encode('utf-8')).hexdigest()[:16] + ext
            path = os.path.join(media_dir, name)
            if not os.path.isfile(path):
                resp = requests.get(url, timeout=15)
                resp.raise_for_status()
                with open(path, 'wb') as f:
                    f.write(resp.content)
            return path
        except Exception as e:
            logger.warning(f"下载远程图片失败: {url} - {e}")
            return None
    
    def _add_paragraph(self, doc, text):
        """添加段落"""
        # 处理粗体和斜体
        text = self._format_text(text)
        p = doc.add_paragraph()
        
        # 简单的文本格式化
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                p.add_run(part)
    
    def _format_text(self, text):
        """格式化文本"""
        # 这里可以添加更多的文本格式化逻辑
        return text
    
    def convert_file(self, markdown_file_path: str, output_filename: Optional[str] = None) -> str:
        """
        转换Markdown文件到Word文档
        
        Args:
            markdown_file_path: Markdown文件路径
            output_filename: 输出文件名
            
        Returns:
            生成的Word文件路径
        """
        try:
            # 读取Markdown文件
            with open(markdown_file_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            # 转换到Word
            return self.markdown_to_word(markdown_content, output_filename)
            
        except Exception as e:
            logger.error(f"转换Markdown文件时出错: {str(e)}")
            raise Exception(f"转换Markdown文件失败: {str(e)}")

    def word_to_markdown(self, docx_path: str) -> str:
        """
        将Word文档（.docx）转换为简化的Markdown文本。
        说明：为满足"在原文件上进行修改"的需求，将Word提取为Markdown，便于用大模型修改后再导出。
        """
        try:
            if not os.path.isfile(docx_path):
                raise FileNotFoundError(f"文件不存在: {docx_path}")
            doc = Document(docx_path)
            return self._extract_content_from_doc(doc)
        except Exception as e:
            logger.error(f"Word 转 Markdown 失败: {str(e)}")
            raise

    def word_to_markdown_from_bytes(self, docx_bytes: bytes) -> str:
        """
        直接从字节数据将Word文档转换为Markdown文本，无需保存临时文件。

        Args:
            docx_bytes: Word文档的字节数据

        Returns:
            转换后的Markdown文本
        """
        try:
            import tempfile
            import os

            # 创建临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(docx_bytes)
                temp_path = temp_file.name

            try:
                # 从临时文件加载文档
                doc = Document(temp_path)
                return self._extract_content_from_doc(doc)
            finally:
                # 确保删除临时文件
                try:
                    os.unlink(temp_path)
                except OSError:
                    pass

        except Exception as e:
            logger.error(f"从字节数据转换Word到Markdown失败: {str(e)}")
            raise

    def _extract_content_from_doc(self, doc) -> str:
        """
        从Document对象中提取所有内容并转换为Markdown

        Args:
            doc: python-docx的Document对象

        Returns:
            转换后的Markdown文本
        """
        try:
            parts = []

            # 提取所有段落
            for paragraph in doc.paragraphs:
                p_md = self._paragraph_to_markdown_full(paragraph)
                if p_md and p_md.strip():
                    parts.append(p_md)

            # 提取所有表格
            for table in doc.tables:
                table_md = self._table_to_markdown_full(table)
                if table_md and table_md.strip():
                    parts.append(table_md)

            # 提取文档中的图片（从XML中）
            for block in doc.element.body:
                if block.tag.endswith('}p'):
                    image_md = self._extract_images_from_xml(block)
                    if image_md and image_md.strip():
                        parts.append(image_md)

            return "\n\n".join(parts)
        except Exception as e:
            logger.error(f"从Document对象提取内容失败: {str(e)}")
            raise

    def word_to_markdown_pandoc(self, docx_path: str) -> str:
        """使用 Pandoc 高保真将 Word 转换为 Markdown。"""
        try:
            import pypandoc
            logger.info("使用 Pandoc 将 Word 提取为 Markdown...")
            if not os.path.isfile(docx_path):
                raise FileNotFoundError(f"文件不存在: {docx_path}")
            # 将媒体资源提取到 output_dir/media，并将链接指向相对路径 media/xxx
            media_dir_name = 'media'
            media_dir_abs = os.path.join(self.output_dir, media_dir_name)
            os.makedirs(media_dir_abs, exist_ok=True)
            md = pypandoc.convert_file(
                docx_path,
                'markdown',
                extra_args=[f"--extract-media={media_dir_abs}"]
            )
            # 规范化为相对 media 路径，避免绝对路径导致资源找不到
            normalized = md.replace(media_dir_abs.replace('\\','/'), media_dir_name)
            return normalized
        except Exception as e:
            logger.error(f"Pandoc 提取失败: {str(e)}，回退到简化提取")
            return self.word_to_markdown(docx_path)

    def word_to_markdown_pandoc_from_bytes(self, docx_bytes: bytes) -> str:
        """
        使用Pandoc直接从字节数据将Word文档转换为Markdown，无需保存临时文件。

        Args:
            docx_bytes: Word文档的字节数据

        Returns:
            转换后的Markdown文本
        """
        try:
            import pypandoc
            import tempfile
            import os

            logger.info("使用 Pandoc 从字节数据提取 Word 为 Markdown...")

            # 创建临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(docx_bytes)
                temp_path = temp_file.name

            try:
                # 将媒体资源提取到 output_dir/media，并将链接指向相对路径 media/xxx
                media_dir_name = 'media'
                media_dir_abs = os.path.join(self.output_dir, media_dir_name)
                os.makedirs(media_dir_abs, exist_ok=True)

                md = pypandoc.convert_file(
                    temp_path,
                    'markdown',
                    extra_args=[f"--extract-media={media_dir_abs}"]
                )

                # 规范化为相对 media 路径，避免绝对路径导致资源找不到
                normalized = md.replace(media_dir_abs.replace('\\','/'), media_dir_name)
                return normalized
            finally:
                # 确保删除临时文件
                try:
                    os.unlink(temp_path)
                except OSError:
                    pass

        except Exception as e:
            logger.error(f"Pandoc 从字节数据提取失败: {str(e)}，回退到简化提取")
            return self.word_to_markdown_from_bytes(docx_bytes)

    def has_pandoc(self) -> bool:
        try:
            import pypandoc
            pypandoc.get_pandoc_version()
            return True
        except Exception:
            return False

    def _paragraph_to_markdown_full(self, paragraph) -> str:
        """完整处理段落转换为Markdown"""
        try:
            # 检测图片：若段落包含图片，导出图片并返回图片Markdown
            image_md = self._extract_images_from_paragraph(paragraph)
            if image_md:
                return image_md

            text = paragraph.text or ""
            if text.strip() == "":
                return ""

            # 标题映射
            style_name = (getattr(paragraph.style, 'name', '') or "").lower()
            for level in range(1, 6):
                if f"heading {level}" in style_name:
                    return f"{'#' * level} {text}"

            # 处理列表项
            level = 0
            is_list = False
            is_ordered = False
            try:
                pPr = paragraph._p.pPr
                if pPr is not None and pPr.numPr is not None:
                    is_list = True
                    ilvl = pPr.numPr.ilvl
                    if ilvl is not None and ilvl.val is not None:
                        level = int(ilvl.val)
                    # 检查是否为有序列表
                    numId = pPr.numPr.numId
                    if numId is not None:
                        # 简化判断：通常numId为1是无序列表，其他是有序列表
                        is_ordered = str(numId.val) != "1"
            except Exception:
                pass

            if is_list:
                indent = "  " * max(0, level)
                prefix = "1. " if is_ordered else "- "
                return f"{indent}{prefix}{text}"

            # 处理引用样式
            if "quote" in style_name:
                return f"> {text}"

            return text
        except Exception as e:
            logger.warning(f"段落转换失败: {e}")
            return paragraph.text if paragraph.text else ""

    def _paragraph_to_markdown(self, p_elm) -> str:
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn
        try:
            paragraph = Paragraph(p_elm, None)
            return self._paragraph_to_markdown_full(paragraph)
        except Exception:
            return ""

    def _table_to_markdown_full(self, table) -> str:
        """完整处理表格转换为Markdown"""
        try:
            rows = table.rows
            if not rows:
                return ""

            # 构建表格Markdown
            md_lines = []
            for row_idx, row in enumerate(rows):
                cells = []
                for cell in row.cells:
                    cell_text = self._clean_cell_text(cell.text)
                    cells.append(cell_text)

                if row_idx == 0:
                    # 表头
                    md_lines.append("| " + " | ".join(cells) + " |")
                    # 分隔线
                    separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                    md_lines.append(separator)
                else:
                    # 数据行
                    md_lines.append("| " + " | ".join(cells) + " |")

            return "\n".join(md_lines)
        except Exception as e:
            logger.warning(f"表格转换失败: {e}")
            return ""

    def _table_to_markdown(self, tbl_elm) -> str:
        from docx.table import Table
        try:
            table = Table(tbl_elm, None)
            return self._table_to_markdown_full(table)
        except Exception:
            return ""

    def _clean_cell_text(self, text: str) -> str:
        """清理表格单元格文本"""
        if not text:
            return ""
        # 移除多余的换行符和空格
        return " ".join(text.split()).strip()

    def _extract_images_from_xml(self, block) -> str:
        """从XML块中提取图片"""
        try:
            from docx.text.paragraph import Paragraph
            paragraph = Paragraph(block, None)
            return self._extract_images_from_paragraph(paragraph)
        except Exception:
            return ""

    def _table_to_markdown(self, tbl_elm) -> str:
        from docx.table import Table
        try:
            table = Table(tbl_elm, None)
            rows = table.rows
            if not rows:
                return ""
            # 取第一行为表头
            header_cells = [self._cell_text_with_merge_hint(c) for c in rows[0].cells]
            sep_cells = ["---" for _ in header_cells]
            md_lines = ["| " + " | ".join(header_cells) + " |", "| " + " | ".join(sep_cells) + " |"]
            for r in rows[1:]:
                md_lines.append("| " + " | ".join([self._cell_text_with_merge_hint(c) for c in r.cells]) + " |")
            return "\n".join(md_lines)
        except Exception:
            return ""

    def _cell_text_with_merge_hint(self, cell) -> str:
        """返回单元格文本并附带合并提示（Markdown不支持合并，故添加提示）。"""
        try:
            text = (cell.text or "").strip()
            tcPr = getattr(cell._tc, 'tcPr', None)
            hints = []
            if tcPr is not None:
                gridSpan = getattr(tcPr, 'gridSpan', None)
                if gridSpan is not None and getattr(gridSpan, 'val', None):
                    hints.append(f"colspan={gridSpan.val}")
                vMerge = getattr(tcPr, 'vMerge', None)
                if vMerge is not None:
                    val = getattr(vMerge, 'val', None)
                    if val is None or str(val) == 'continue':
                        hints.append("rowspan=continue")
                    elif str(val) == 'restart':
                        hints.append("rowspan=start")
            if hints:
                return f"{text} <{' '.join(hints)}>"
            return text
        except Exception:
            return (cell.text or "").strip()

    def _extract_images_from_paragraph(self, paragraph) -> str:
        """提取段落中的第一张图片为文件并返回图片Markdown；若无图片则返回空串。"""
        try:
            # 查找 run 中的图片关系
            for run in paragraph.runs:
                drawing_elms = run._element.xpath('.//a:blip')
                if not drawing_elms:
                    continue
                # 取首个图片关系id
                r_embed = drawing_elms[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if not r_embed:
                    continue
                # 定位图片part
                image_part = paragraph.part.related_parts.get(r_embed)
                if image_part is None:
                    continue
                # 保存图片到输出目录
                os.makedirs(self.output_dir, exist_ok=True)
                ext = os.path.splitext(image_part.partname.basename())[1] or '.png'
                from datetime import datetime
                img_name = f"extracted_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}{ext}"
                img_path = os.path.join(self.output_dir, img_name)
                with open(img_path, 'wb') as fp:
                    fp.write(image_part.blob)
                # 返回Markdown图片语法
                return f"![image]({img_name})"
        except Exception as e:
            logger.warning(f"提取图片失败: {e}")
        return ""
